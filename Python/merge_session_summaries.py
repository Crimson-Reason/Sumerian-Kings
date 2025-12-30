#!/usr/bin/env python3
"""Merge multiple Session_Summary_Analysis .docx files into a single combined document.

This script looks for files matching Session_Summary_Analysis* in the Documentation/ folder,
appends them in chronological order, and writes Combined_Session_Summary_Analysis_{date}.docx
"""
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
from datetime import date
import re

DOC_DIR = Path(__file__).resolve().parent.parent / "Documentation"
OUT_TEMPLATE = "Combined_Session_Summary_Analysis_{}.docx"


def iter_block_items(parent):
    """Yield each paragraph or table element from a Document or cell.
    Adapted from python-docx docs/recipes.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)


def copy_paragraph(src_p, dst_doc):
    """Copy paragraph text and basic style to dst_doc as a new paragraph."""
    p = dst_doc.add_paragraph()
    if src_p.style is not None:
        try:
            p.style = src_p.style
        except Exception:
            pass
    for run in src_p.runs:
        r = p.add_run(run.text)
        try:
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.size = run.font.size
            r.font.name = run.font.name
        except Exception:
            pass
    # copy alignment
    try:
        p.alignment = src_p.alignment
    except Exception:
        pass
    return p


def copy_table(src_tbl, dst_doc):
    rows = len(src_tbl.rows)
    cols = len(src_tbl.columns)
    tbl = dst_doc.add_table(rows=rows, cols=cols)
    tbl.style = src_tbl.style
    for i, row in enumerate(src_tbl.rows):
        for j, cell in enumerate(row.cells):
            # join paragraphs in cell
            text = '\n'.join(p.text for p in cell.paragraphs)
            tbl.cell(i, j).text = text
    return tbl


def find_session_files():
    pattern = re.compile(r"Session_Summary_Analysis.*\.docx$")
    files = [f for f in DOC_DIR.iterdir() if f.is_file() and pattern.match(f.name)]
    # sort by modified time (oldest first)
    files.sort(key=lambda p: p.stat().st_mtime)
    return files


def main():
    files = find_session_files()
    if not files:
        print("No session summary files found in Documentation/ (pattern: Session_Summary_Analysis*.docx)")
        return 1
    print(f"Found {len(files)} session files:")
    for f in files:
        print(" -", f.name)

    out_doc = Document()

    for idx, f in enumerate(files, 1):
        src = Document(f)
        # Add source heading and page break if not first
        if idx > 1:
            out_doc.add_page_break()
        heading = out_doc.add_heading(f.name, level=2)
        heading.alignment = None

        # Copy blocks
        for block in iter_block_items(src):
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            if isinstance(block, Paragraph):
                # skip empty paragraph that's only the heading we added
                if block.text.strip():
                    copy_paragraph(block, out_doc)
            elif isinstance(block, Table):
                copy_table(block, out_doc)

    out_name = DOC_DIR / OUT_TEMPLATE.format(date.today().strftime("%Y%m%d"))
    out_doc.save(out_name)
    print(f"Combined document written: {out_name}")
    return 0


if __name__ == '__main__':
    main()
